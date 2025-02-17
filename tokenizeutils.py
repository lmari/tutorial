import os
import numpy as np
from colorama import Fore, Style
from docx import Document
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM
import logging
from transformers import logging as transformers_logging

logging.basicConfig(level=logging.ERROR)
transformers_logging.set_verbosity_error()

def cosine_similarity(a: np.ndarray, b: np.ndarray) -> float:
    '''Calcola il coseno di similitudine ("cosine similarity",
        https://it.wikipedia.org/wiki/Coseno_di_similitudine) tra i due vettori a e b.'''
    return np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b))


def cosine_similarities(a: np.ndarray, bs: list[np.ndarray]) -> list[float]:
    '''Calcola il coseno di similitudine tra il vettore a e ognuno dei vettori bs.'''
    norm_a = np.linalg.norm(a)
    return [np.dot(a, b) / (norm_a * np.linalg.norm(b)) for b in bs]


def distance(a: np.ndarray, b: np.ndarray) -> float:
    '''Calcola la distanza tra i due vettori a e b.'''
    return np.linalg.norm(a - b) # type: ignore


def colorize(tokens, clean=False):
    c = [Fore.RED, Fore.BLUE]
    i = 0
    x = ""
    for token in tokens:
        x += c[i] + " " + (token if token[:2] != "##" or not clean else token[2:])
        i = 1 - i
    x += Style.RESET_ALL
    return x


def read_file(name):
    '''Legge il contenuto di un file docx e restituisce il testo come stringa.'''
    return "\n".join([p.text for p in Document(name).paragraphs])


def read_file2(name):
    '''Legge il contenuto di un file docx e restituisce il testo come lista di paragrafi.'''
    x = [p.text.strip() for p in Document(name).paragraphs]
    x = x[7:]                       # rimuovi l'intestazione
    x = [y for y in x if y != '']   # rimuovi i paragrafi vuoti
    min_len = 100
    y = [x[0]]                      # unisci i paragrafi troppo corti
    for i in range(1, len(x)-1):
        if len(y[-1]) < min_len:
            y[-1] += " " + x[i]
        else:
            y.append(x[i])
    if len(x[-1]) < min_len or len(y[-1]) < min_len:    # unisci l'ultimo paragrafo se è troppo corto
        y[-1] += " " + x[-1]
    else:
        y.append(x[-1])
    return y


def read_file3(name):
    '''Legge il contenuto di un file docx e restituisce il testo come lista di frasi.'''
    paragraphs = [p.text.strip() for p in Document(name).paragraphs]
    paragraphs = paragraphs[7:]                         # rimuovi l'intestazione
    paragraphs = [p for p in paragraphs if p != '']     # rimuovi i paragrafi vuoti
    result = []
    for p in paragraphs:
        sentences = p.split(".")                        # separa i paragrafi in frasi
        sentences = [s for s in sentences if len(s.strip()) > 10]   # rimuovi le frasi vuote o troppo corte
        result.extend(sentences)
    return result


def read_files(metadata, doc_dir):
    '''Legge i file indicati nei metadata e restituisce una lista di tuple (titolo, testo).'''
    texts = []
    for data in metadata:
        doc_text = read_file(os.path.join(doc_dir, data[0]))
        texts.append((data[1], doc_text))
    return texts


def read_files2(metadata, doc_dir):
    '''Legge i file indicati nei metadata e restituisce una lista di tuple (titolo, indice_paragrafo, testo).'''
    texts = []
    for data in metadata:
        par_text = read_file2(os.path.join(doc_dir, data[0]))
        for i, text in enumerate(par_text):
            texts.append((data[0], i, text))
    return texts


def read_files3(metadata, doc_dir):
    '''Legge i file indicati nei metadata e restituisce una lista di tuple (titolo, indice_frase, testo).'''
    texts = []
    for data in metadata:
        par_text = read_file3(os.path.join(doc_dir, data[0]))
        for i, text in enumerate(par_text):
            texts.append((data[0], i, text))
    return texts


class Model():

    def __init__(self, model_id, cased=False):
        self.tokenizer = AutoTokenizer.from_pretrained(model_id)
        self.embedder = AutoModelForCausalLM.from_pretrained(model_id, is_decoder=False)
        self.cased = cased
        self.embedding_layer = self.embedder.get_input_embeddings()             # il layer di embedding del modello
        self.vocab_embeddings = self.embedding_layer.weight.detach().numpy()    # gli embedding dell'intero vocabolario
        self.vocab_size = self.tokenizer.vocab_size
        self.embedding_dim = self.embedding_layer.embedding_dim

    
    def c(self, token: str) -> str:
        '''Restituisce il token in forma cased o uncased in accordo al modello'''
        return token if self.cased else token.lower()


    def token_in_vocab(self, token: str) -> bool:
        '''Restituisce True se il token è presente nel vocabolario del modello, False altrimenti.'''
        return self.c(token) in self.tokenizer.get_vocab()


    def token_to_id(self, token: str) -> int:
        '''Dato un token, ne restituisce l'id nel vocabolario del modello.'''
        return self.tokenizer.convert_tokens_to_ids(self.c(token)) # type: ignore


    def token_to_embedding(self, token: str) -> np.ndarray:
        '''Dato un token, ne restituisce l'embedding.'''
        if not self.token_in_vocab(self.c(token)):
            return np.array([]) # f"Il token '{self.c(token)}' non è presente nel vocabolario del modello."
        token_id = self.tokenizer.convert_tokens_to_ids(self.c(token))
        return self.vocab_embeddings[token_id]
    

    def similarity(self, token1: str, token2: str) -> float:
        '''Calcola la similarità tra due token.'''
        if not self.token_in_vocab(self.c(token1)):
            return 0 # f"Il token '{self.c(token1)}' non è presente nel vocabolario del modello."
        if not self.token_in_vocab(self.c(token2)):
            return 0 # f"Il token '{self.c(token2)}' non è presente nel vocabolario del modello."
        token1_embedding = self.token_to_embedding(token1)
        token2_embedding = self.token_to_embedding(token2)
        return cosine_similarity(token1_embedding, token2_embedding)


    def rough_embed(self, text):
        '''Dato un testo, ne restituisce l'embedding,
        come media degli embedding dei suoi token.'''
        tokens = np.array(self.tokenizer.tokenize(text))
        ids = np.array(self.tokenizer.convert_tokens_to_ids(tokens)) # type: ignore
        embeddings = self.vocab_embeddings[ids]
        return np.mean(embeddings, axis=0)


    def better_embed(self, text):
        '''Dato un testo, ne restituisce l'embedding,
        calcolato da un transformer.'''
        encoded_input = self.tokenizer(text, return_tensors='pt', padding=True, truncation=True, max_length=512)
        with torch.no_grad():
            model_output = self.embedder(**encoded_input)
        return model_output.pooler_output.squeeze()


    def most_similar(self, token: str, top_n: int=1, filter: bool=True) -> list[str]:
        '''Calcola la lista dei token più simili al token dato.
        Se filter=True, sono esclusi i token che sono simili al token dato,
        che iniziano con '##' o che sono numeri.'''
        if not self.token_in_vocab(self.c(token)):
            return  [] # f"Il token '{self.c(token)}' non è presente nel vocabolario del modello."
        token = self.c(token)
        token_embedding = self.token_to_embedding(token)
        similarities = cosine_similarities(token_embedding, self.vocab_embeddings)
        most_similar_ids = reversed(np.argsort(similarities))
        i = 0
        result = []
        for id in most_similar_ids:
            candidate_token = self.tokenizer.convert_ids_to_tokens([id])[0]
            if not filter or (candidate_token.lower() != token.lower() and not candidate_token.startswith('##') and not candidate_token.isdigit()):
                result.append((candidate_token, round(float(similarities[id]),2)))
                i += 1
            if i == top_n: break
        return result


    def combine_meanings(self, positive_tokens: list[str], negative_tokens: list=[str], top_n: int=1, filter:bool=True) -> list[str]:
        '''Calcola gli n token più simili ai token dati'''
        for token in positive_tokens + negative_tokens:
            if not self.token_in_vocab(self.c(token)):
                return [] # f"Il token '{self.c(token)}' non è presente nel vocabolario del modello."
        token_embeddings = np.zeros(self.embedding_layer.embedding_dim)
        for token in positive_tokens:
            token_embeddings += self.token_to_embedding(token) # type: ignore
        for token in negative_tokens:
            token_embeddings -= self.token_to_embedding(token) # type: ignore
        similarities = cosine_similarities(token_embeddings, self.vocab_embeddings)
        reference_tokens = list(map(str.lower, positive_tokens + negative_tokens))
        most_similar_ids = reversed(np.argsort(similarities))
        i = 0
        result = []
        for id in most_similar_ids:
            candidate_token = self.tokenizer.convert_ids_to_tokens([id])[0]
            if not filter or not self._filter_out(candidate_token, reference_tokens): #(not candidate_token.lower() in reference_tokens and not candidate_token.startswith('##') and not candidate_token.isdigit()):
                result.append((candidate_token, round(float(similarities[id]),2)))
                i += 1
            if i == top_n: break
        return result


    def mean_meanings(self, token1: str, token2: str, top_n: int=1, filter:bool=True) -> list[str]:
        if not self.token_in_vocab(self.c(token1)):
            return  [] # f"Il token '{self.c(token1)}' non è presente nel vocabolario del modello."
        if not self.token_in_vocab(self.c(token2)):
            return  [] # f"Il token '{self.c(token2)}' non è presente nel vocabolario del modello."
        token1 = self.c(token1)
        token1_embedding = self.token_to_embedding(token1)
        token2 = self.c(token2)
        token2_embedding = self.token_to_embedding(token2)
        mean_embedding = (token1_embedding + token2_embedding) / 2
        similarities = cosine_similarities(mean_embedding, self.vocab_embeddings)
        most_similar_ids = reversed(np.argsort(similarities))
        i = 0
        result = []
        for id in most_similar_ids:
            candidate_token = self.tokenizer.convert_ids_to_tokens([id])[0]
            if not filter or not self._filter_out(candidate_token, [token1, token2]):
                result.append((candidate_token, round(similarities[id],2)))
                i += 1
            if i == top_n: break
        return result
    

    def meaning_distance(self, token1: str, token2: str) -> float:
        if not self.token_in_vocab(self.c(token1)):
            return 0
        if not self.token_in_vocab(self.c(token2)):
            return 0
        return distance(self.token_to_embedding(self.c(token1)), self.token_to_embedding(self.c(token2)))
    

    def _filter_out(self, token: str, reference_tokens: list[str]) -> bool:
        '''Restituisce True se il token deve essere filtrato.'''
        reference_tokens = list(map(str.lower, reference_tokens))
        return token.lower() in reference_tokens or token.startswith('##') or token.isdigit()
